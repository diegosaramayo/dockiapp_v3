import docx

class DocxParser:
    def parse(self, file_path):
        fields = []
        doc = docx.Document(file_path)

        field_idx = 1
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            if any(term in text.lower() for term in ["declaración", "declaracion", "jurada", "terminos", "condiciones"]):
                if len(text) < 60:
                    fields.append({
                        "id": f"docx_field_{field_idx}",
                        "type": "terms_header",
                        "text": text if text.endswith(":") else f"{text}:"
                    })
                else:
                    fields.append({
                        "id": f"docx_field_{field_idx}",
                        "type": "terms_box",
                        "text": text
                    })
            elif ":" in text and len(text) < 100:
                parts = text.split(":", 1)
                label = parts[0].strip() + ":"
                val = parts[1].strip()

                fields.append({
                    "id": f"docx_field_{field_idx}",
                    "type": "input",
                    "label": label,
                    "placeholder": val if val else "XXXX",
                    "value": ""
                })
            elif text.startswith("[ ]") or text.startswith("[x]") or text.startswith("☐") or "encuadrado" in text.lower():
                clean_label = text.replace("[ ]", "").replace("[x]", "").replace("☐", "").strip()
                fields.append({
                    "id": f"docx_field_{field_idx}",
                    "type": "check_options",
                    "label": clean_label if clean_label else text,
                    "is_checked": None
                })
            else:
                if len(text) > 100:
                    fields.append({
                        "id": f"docx_field_{field_idx}",
                        "type": "terms_box",
                        "text": text
                    })
                else:
                    fields.append({
                        "id": f"docx_field_{field_idx}",
                        "type": "input",
                        "label": text if text.endswith(":") else f"{text}:",
                        "placeholder": "XXXX",
                        "value": ""
                    })
            field_idx += 1

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if len(cells) >= 2:
                    label = cells[0] if cells[0].endswith(":") else f"{cells[0]}:"
                    fields.append({
                        "id": f"docx_field_{field_idx}",
                        "type": "input",
                        "label": label,
                        "placeholder": cells[1],
                        "value": ""
                    })
                    field_idx += 1

        if not fields:
            fields.append({
                "id": "docx_default_1",
                "type": "input",
                "label": "Campo Extraído DOCX:",
                "placeholder": "XXXX",
                "value": ""
            })
        return fields
